import streamlit as st
import google.generativeai as genai
from PIL import Image
import pytesseract
from docx import Document
from io import BytesIO
import datetime

# ====== Configure Gemini ======
genai.configure(api_key="AQ.Ab8RN6JRxoGnxSlZ1h60io08XGdGKXm-wWG2hFftiWEfhafk2A")
model = genai.GenerativeModel("gemini-2.0-flash")

# ====== Set Tesseract path ======
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

st.set_page_config(page_title="Medical Summarizer", layout="centered")
st.title("🩺 Medical Text Summarizer & Q&A Generator")

option = st.radio("Choose input method:", ["📝 Paste Text", "🖼️ Upload Image"])

user_text = ""

if option == "📝 Paste Text":
    user_text = st.text_area("Paste your medical text here:", height=200)

else:
    uploaded_file = st.file_uploader("Upload image (JPG, PNG)", type=["jpg", "jpeg", "png"])
    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption="Uploaded Image", width=300)
        with st.spinner("📖 Extracting English text..."):
            try:
                user_text = pytesseract.image_to_string(image, lang='eng')
            except Exception as e:
                st.error(f"OCR Error: {e}")
                user_text = ""
        if user_text.strip():
            st.success("✅ Text extracted!")
            st.text_area("Extracted text (edit if needed):", user_text, height=150)
        else:
            st.warning("⚠️ No English text found.")

# ====== Analysis Button ======
if st.button("🚀 Analyze"):
    if user_text.strip():
        with st.spinner("Processing with AI..."):
            prompt = f"""
            Summarize the following medical text into clear bullet points, 
            then generate 10 Q&As based strictly on the text. 
            Do not omit or add any information. Output in English.
            Text: {user_text}
            """
            response = model.generate_content(prompt)
            result_text = response.text
            
            st.success("✅ Analysis Done!")
            st.write(result_text)

            # ====== NEW: Word Download Button ======
            # Create Word document in memory
            doc = Document()
            doc.add_heading('Medical Text Analysis Report', 0)
            doc.add_heading('Original Text:', level=1)
            doc.add_paragraph(user_text)
            
            doc.add_heading('Analysis Result (Summary & Q&A):', level=1)
            # Split the result into paragraphs for better formatting
            for line in result_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            
            # Add timestamp
            doc.add_paragraph(f"\nGenerated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

            # Save to bytes
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Download button
            st.download_button(
                label="📥 Download as Word Document (.docx)",
                data=doc_io,
                file_name="Medical_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml"
            )
    else:
        st.warning("Please provide text or image first.")